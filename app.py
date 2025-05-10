# Ensure you've installed the necessary libraries manually before running the script
# In your terminal, run: 
# pip install streamlit transformers fpdf python-docx

import streamlit as st
from transformers import pipeline
from fpdf import FPDF
from docx import Document

# 🛠 Custom CSS for UI styling
st.markdown("""
    <style>
    body {
        background-color: #f0f2f6;
    }
    .main {
        background-color: #ffffff;
        padding: 20px;
        border-radius: 10px;
    }
    h1 {
        color: #ff4b4b;
        text-align: center;
        font-family: 'Trebuchet MS', sans-serif;
    }
    </style>
""", unsafe_allow_html=True)

# 🖼 Logo and Title
st.image("https://upload.wikimedia.org/wikipedia/commons/4/4f/Iconic_image.png", width=100)
st.title("📝 Smart Article Summarizer")

# 🛠 Upload or Paste Text
option = st.radio("Choose input method:", ("Upload a text file", "Paste text"))

input_text = ""

if option == "Upload a text file":
    uploaded_file = st.file_uploader("Upload your text file", type=["txt"])
    if uploaded_file is not None:
        input_text = uploaded_file.read().decode("utf-8")
else:
    input_text = st.text_area("Paste your article here:", height=300)

# 🛠 Sidebar Settings for Summary
st.sidebar.subheader("Summary Settings")

language = st.sidebar.selectbox('Select Language:', ('English', 'French', 'Hindi'))
max_length = st.sidebar.slider('Maximum summary length:', 50, 500, 150)
min_length = st.sidebar.slider('Minimum summary length:', 10, 100, 50)

# 🛠 Load Summarizer Model
if language in ['French', 'Hindi']:
    summarizer = pipeline("summarization", model="csebuetnlp/mT5_multilingual_XLSum")
else:
    summarizer = pipeline("summarization")

# 🛠 Generate Summary
summary_text = ""

if st.button("✨ Generate Summary"):
    if input_text.strip() != "":
        with st.spinner('Summarizing... Please wait ⏳'):
            summary = summarizer(input_text, max_length=max_length, min_length=min_length, do_sample=False)
            summary_text = summary[0]['summary_text']
        st.success('Summary generated successfully!')
        st.subheader("Summary Output:")
        st.write(summary_text)
    else:
        st.error("Please upload a file or paste some text!")

# 🛠 Download Options (PDF and Word)
def save_as_pdf(summary_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary_text)
    pdf.output("summary.pdf")

def save_as_docx(summary_text):
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary_text)
    doc.save("summary.docx")

if summary_text:
    st.subheader("Download Summary:")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("📄 Download as PDF"):
            save_as_pdf(summary_text)
            with open("summary.pdf", "rb") as f:
                st.download_button(label="Click to Download PDF", data=f, file_name="summary.pdf", mime='application/pdf')

    with col2:
        if st.button("📝 Download as Word"):
            save_as_docx(summary_text)
            with open("summary.docx", "rb") as f:
                st.download_button(label="Click to Download Word", data=f, file_name="summary.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
